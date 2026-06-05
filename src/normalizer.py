import re
from dataclasses import dataclass
from typing import Optional, Dict, List, Tuple, Any

# =========================
# Limpieza base (lo tuyo)
# =========================
RE_DATOS_PERSONALES = re.compile(
    r"DATOS\s+PERSONALES[\s\S]*?FORMACI[ÓO]N\s+ACAD[ÉE]MICA",
    re.IGNORECASE
)
RE_RUBRO_BASURA = re.compile(
    r"^\s*TECNOLOG[IÍ]A\s+E\s+INNOVACI[ÓO]N\s*$",
    re.IGNORECASE | re.MULTILINE
)
RE_NULLS = re.compile(r"\bnull(?:\s*\(ed\))?\b", re.IGNORECASE)
RE_MANY_NEWLINES = re.compile(r"\n{3,}")
RE_SPACES = re.compile(r"[ \t]+")
RE_CVAR_BOILERPLATE_LINE = re.compile(
    r"(?:CVar\s+ES\s+UNA\s+INICIATIVA|MINISTERIO\s+DE\s+CIENCIA|"
    r"TECNOLOG[IÍ]A\s+E\s+INNOVACI[ÓO]N|Fecha\s+de\s+generaci[oó]n)",
    re.IGNORECASE,
)
RE_PAGE_NUMBER_LINE = re.compile(r"^\s*\d{1,3}\s*$")
RE_CVAR_NAME_REPEAT = re.compile(
    r"^[A-ZÁÉÍÓÚÜÑ][A-ZÁÉÍÓÚÜÑ\s.'-]+,\s*[A-ZÁÉÍÓÚÜÑ][A-ZÁÉÍÓÚÜÑ\s.'-]+$",
    re.MULTILINE,
)

@dataclass
class NormalizeOptions:
    remove_personal_data: bool = True
    remove_rubros_basura: bool = True
    remove_nulls: bool = True
    collapse_spaces: bool = True

def normalize_text(raw_text: str, opts: Optional[NormalizeOptions] = None) -> str:
    opts = opts or NormalizeOptions()
    text = (raw_text or "").replace("\r\n", "\n").replace("\r", "\n")

    if opts.remove_personal_data:
        # deja el encabezado para que exista el punto de anclaje
        text = RE_DATOS_PERSONALES.sub("FORMACIÓN ACADÉMICA\n", text)

    if opts.remove_rubros_basura:
        text = RE_RUBRO_BASURA.sub("", text)

    if opts.remove_nulls:
        text = RE_NULLS.sub("", text)

    if opts.collapse_spaces:
        # normaliza espacios por línea (sin pegar líneas)
        text = "\n".join(RE_SPACES.sub(" ", ln).strip() for ln in text.split("\n"))

    cleaned_lines = []
    for ln in text.split("\n"):
        stripped = ln.strip()
        if not stripped:
            cleaned_lines.append("")
            continue
        if RE_CVAR_BOILERPLATE_LINE.search(stripped):
            continue
        if RE_PAGE_NUMBER_LINE.match(stripped):
            continue
        if RE_CVAR_NAME_REPEAT.match(stripped):
            continue
        cleaned_lines.append(stripped)
    text = "\n".join(cleaned_lines)

    text = RE_MANY_NEWLINES.sub("\n\n", text).strip() + "\n"
    return text


# =========================
# PDF -> texto
# =========================
def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    import io
    import pdfplumber
    parts = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts)


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    import io
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_from_doc_bytes(doc_bytes: bytes) -> str:
    """Extrae texto de .doc (Word antiguo) usando textutil en macOS."""
    import os
    import subprocess
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(doc_bytes)
        path = tmp.name
    try:
        proc = subprocess.run(
            ["textutil", "-stdout", "-convert", "txt", path],
            capture_output=True,
            text=True,
            check=False,
        )
        if proc.returncode == 0 and proc.stdout.strip():
            return proc.stdout
        raise ValueError(proc.stderr.strip() or "No se pudo convertir el archivo .doc")
    finally:
        os.unlink(path)


# ==========================================================
# 1) Secciones: extracción por BLOQUES
# ==========================================================
# Encabezados típicos de CVar (variantes + normalización)
SECTION_ALIASES: List[Tuple[str, str]] = [
    # canonical, regex alias
    ("FORMACIÓN ACADÉMICA", r"FORMACI[ÓO]N\s+ACAD[ÉE]MICA(?:\s+Y\s+COMPLEMENTARIA)?"),
    ("FORMACIÓN DE RECURSOS HUMANOS", r"FORMACI[ÓO]N\s+DE\s+RECURSOS\s+HUMANOS|RECURSOS\s+HUMANOS|RRHH"),
    ("PRODUCCIÓN CIENTÍFICA", r"PRODUCCI[ÓO]N\s+CIENT[IÍ]FICA|PUBLICACIONES"),
    ("FINANCIAMIENTO CIENTÍFICO Y TECNOLÓGICO", r"FINANCIAMIENTO|PROYECTOS\s+DE\s+I\+D|PROYECTOS\s+I\+D"),
    ("EXTENSIÓN", r"EXTENSI[ÓO]N"),
    ("EVALUACIÓN / GESTIÓN EDITORIAL", r"EVALUACI[ÓO]N|GESTI[ÓO]N\s+EDITORIAL|JURADO"),
    ("PARTICIPACIÓN EN EVENTOS CyT", r"EVENTOS|REUNIONES\s+CIENT[IÍ]FICAS|CONGRESOS|JORNADAS"),
    ("ANTECEDENTES EN CyT", r"ANTECEDENTES"),
    ("IDIOMAS", r"IDIOMAS"),
    ("CURSOS Y CAPACITACIONES", r"CURSOS|CAPACITACIONES|FORMACI[ÓO]N\s+COMPLEMENTARIA"),
]

# patrón “encabezado” = debe aparecer como línea (o casi)
def _build_section_header_regex() -> re.Pattern:
    parts = []
    for _, alias in SECTION_ALIASES:
        parts.append(f"(?:{alias})")
    # ancla a inicio de línea con posibles espacios
    pat = r"(?im)^\s*(%s)\s*$" % "|".join(parts)
    return re.compile(pat)

RE_SECTION_HEADER = _build_section_header_regex()

def _canonical_section_name(found_header_line: str) -> str:
    h = (found_header_line or "").strip().upper()
    for canonical, alias in SECTION_ALIASES:
        if re.fullmatch(alias, h, flags=re.IGNORECASE):
            return canonical
    # fallback suave: intenta match parcial
    for canonical, alias in SECTION_ALIASES:
        if re.search(alias, h, flags=re.IGNORECASE):
            return canonical
    return h.strip()

def extract_sections(clean_text: str) -> Dict[str, str]:
    """
    Devuelve dict {seccion_canonica: bloque_texto} usando cortes por encabezados.
    Si no encuentra encabezados, devuelve {"TEXTO_COMPLETO": clean_text}.
    """
    txt = clean_text or ""
    matches = list(RE_SECTION_HEADER.finditer(txt))
    if not matches:
        return {"TEXTO_COMPLETO": txt.strip()}

    # construir rangos
    spans: List[Tuple[str, int, int]] = []
    for i, m in enumerate(matches):
        start = m.start()
        header_line = m.group(1) if m.group(1) else m.group(0)
        canonical = _canonical_section_name(header_line)

        end = matches[i + 1].start() if i + 1 < len(matches) else len(txt)
        block = txt[m.end():end].strip()
        spans.append((canonical, start, end))

    # armar dict manteniendo el primer bloque de cada sección (si repite, concatena)
    out: Dict[str, str] = {}
    for i, m in enumerate(matches):
        header_line = m.group(1) if m.group(1) else m.group(0)
        canonical = _canonical_section_name(header_line)
        end = matches[i + 1].start() if i + 1 < len(matches) else len(txt)
        block = txt[m.end():end].strip()

        if canonical in out:
            # concatena si aparece nuevamente
            out[canonical] = (out[canonical] + "\n\n" + block).strip()
        else:
            out[canonical] = block

    return out


# ==========================================================
# 2) Formación Académica: parse robusto por “tipos”
#    (para evitar listas infinitas de títulos)
# ==========================================================
RE_IN_PROGRESS = re.compile(r"\b(Actualidad|En\s+curso|Cursando|Actualmente|Vigente|Hasta\s+la\s+actualidad|A\s+la\s+fecha)\b", re.IGNORECASE)
RE_FINISH = re.compile(r"A[nñ]o\s+de\s+(finalizaci[oó]n|obtenci[oó]n|graduaci[oó]n)\s*:\s*([0-3]?\d\s*[/\-]\s*\d{4}|\d{4})", re.IGNORECASE)

def _split_entries_by_blanklines(block: str) -> List[str]:
    if not block:
        return []
    # separa por doble salto
    parts = re.split(r"\n\s*\n+", block.strip())
    return [p.strip() for p in parts if p.strip()]

def _is_completed(entry: str) -> bool:
    if RE_IN_PROGRESS.search(entry):
        return False
    return bool(RE_FINISH.search(entry) or re.search(r"\b(finalizad[oa]|egresad[oa]|graduad[oa]|t[ií]tulo\s+obtenido|completo)\b", entry, re.IGNORECASE))

def _first_nonempty_line(entry: str) -> str:
    for ln in entry.split("\n"):
        ln = ln.strip()
        if ln and ln.lower() != "null":
            return ln
    return ""

def _guess_formacion_type(entry: str) -> str:
    e = entry.lower()

    # Estancias (las querés tratar como “Estancia”)
    if "estancia" in e and ("i+d" in e or "i + d" in e or "investig" in e):
        return "ESTANCIA"

    # Diplomaturas (a veces las ponen como “Especialización en ‘DIPLOMATURA…’”)
    if "diplomatura" in e:
        return "DIPLOMATURA"

    # Doctorado / Doctor
    if re.search(r"\bdoctorad[oa]\b|\bdoctor\s+en\b|\bdoctora\b|\bdoctor\b", e):
        return "DOCTORADO"

    # Maestría / Magíster / Máster
    if re.search(r"\bmaestr[ií]a\b|\bmag[ií]ster\b|\bmaster\b|\bmáster\b", e):
        return "MAESTRÍA"

    # Especialización / Especialista
    if re.search(r"\bespecializaci[oó]n\b|\bespecialista\b", e):
        return "ESPECIALIZACIÓN"

    # Profesorado
    if re.search(r"\bprofesorado\b|\bprofesor\s+universitario\b|\bprofesor\s+en\b", e):
        return "PROFESORADO"

    # Grado (sin lista infinita): regla práctica
    # Si es 1ª línea “Xxxxxxx” y luego “FACULTAD/UNIVERSIDAD” + “Año de finalización”
    # o si contiene "licenciatura/ingenier/abogad/contador/medic/bioquimic/farmac" etc.
    if re.search(r"\b(licenciatura|licenciad|ingenier|abogad|contador|m[eé]dic|bioqu[ií]mic|farmac[eé]utic|arquitect|odont[oó]log)\b", e):
        return "GRADO"

    # Heurística de grado por formato:
    if RE_FINISH.search(entry) and re.search(r"\b(UNIVERSIDAD|FACULTAD|INSTITUTO)\b", entry, re.IGNORECASE):
        # si no es posgrado explícito, lo tratamos como grado
        if not re.search(r"\bdoctorad|\bmaestr|\bmag[ií]ster|\bespecializ", e):
            return "GRADO"

    return "OTRO"

def parse_formacion_academica(block: str) -> Dict[str, List[Dict[str, Any]]]:
    """
    Devuelve estructura:
    {
      "DOCTORADO": [{titulo, texto, finalizado}],
      "MAESTRÍA": [...],
      "ESPECIALIZACIÓN": [...],
      "GRADO": [...],
      "DIPLOMATURA": [...],
      "ESTANCIA": [...],
      "PROFESORADO": [...],
      "OTRO": [...]
    }
    """
    entries = _split_entries_by_blanklines(block)
    out: Dict[str, List[Dict[str, Any]]] = {k: [] for k in [
        "DOCTORADO", "MAESTRÍA", "ESPECIALIZACIÓN", "GRADO",
        "DIPLOMATURA", "ESTANCIA", "PROFESORADO", "OTRO"
    ]}

    for ent in entries:
        if not ent.strip():
            continue
        t = _guess_formacion_type(ent)
        out[t].append({
            "titulo": _first_nonempty_line(ent),
            "finalizado": _is_completed(ent),
            "texto": ent.strip()
        })
    return out


# ==========================================================
# 3) DOCX estructurado (por secciones y sub-bloques)
# ==========================================================
def build_docx_bytes(
    clean_text: str,
    sections: Optional[Dict[str, str]] = None,
    parsed_formacion: Optional[Dict[str, List[Dict[str, Any]]]] = None
) -> bytes:
    import io
    from docx import Document

    doc = Document()

    # Título simple
    doc.add_heading("CVar normalizado", level=1)

    sections = sections or extract_sections(clean_text)

    # Formación académica parseada si existe
    if "FORMACIÓN ACADÉMICA" in sections:
        doc.add_heading("FORMACIÓN ACADÉMICA", level=2)
        block = sections.get("FORMACIÓN ACADÉMICA", "").strip()
        pf = parsed_formacion or parse_formacion_academica(block)

        # orden de presentación
        order = ["DOCTORADO", "MAESTRÍA", "ESPECIALIZACIÓN", "GRADO", "PROFESORADO", "DIPLOMATURA", "ESTANCIA", "OTRO"]
        for k in order:
            items = pf.get(k, [])
            if not items:
                continue
            doc.add_heading(k, level=3)
            for it in items:
                # una línea “limpia” para leer rápido
                titulo = it.get("titulo", "").strip()
                fin = "FINALIZADO" if it.get("finalizado") else "NO FINALIZADO / EN CURSO"
                p = doc.add_paragraph(f"- {titulo} — {fin}")
                # cuerpo completo abajo (para repo 2 o auditoría)
                doc.add_paragraph(it.get("texto", "").strip())

        # Evitar duplicar el bloque bruto de formación
        used_formacion = True
    else:
        used_formacion = False

    # Resto de secciones (sin parse profundo por ahora)
    for sec, block in sections.items():
        if sec == "FORMACIÓN ACADÉMICA":
            continue
        if sec == "TEXTO_COMPLETO" and not used_formacion:
            doc.add_heading("TEXTO COMPLETO", level=2)
        else:
            doc.add_heading(sec, level=2)

        if not block.strip():
            doc.add_paragraph("(Sin contenido detectado)")
            continue

        # escribir bloque preservando saltos
        for para in block.split("\n"):
            para = para.strip()
            if not para:
                continue
            doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ==========================================================
# 4) Helper para generar también JSON estructurado (opcional)
# ==========================================================
def build_structured_output(clean_text: str) -> Dict[str, Any]:
    sections = extract_sections(clean_text)
    out: Dict[str, Any] = {"sections": sections}

    if "FORMACIÓN ACADÉMICA" in sections:
        out["parsed"] = {
            "FORMACIÓN ACADÉMICA": parse_formacion_academica(sections.get("FORMACIÓN ACADÉMICA", ""))
        }
    return out
