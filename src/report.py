import io
from typing import Any, Dict, List, Optional

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from section_caps import (
    allocate_section_item_caps,
    section_effective_max,
    section_uses_shared_pool,
)

CATEGORY_LABELS = {
    "I": "Investigador Superior",
    "II": "Investigador Principal",
    "III": "Investigador Independiente",
    "IV": "Investigador Adjunto",
    "V": "Investigador Asistente",
    "VI": "Becario de Iniciación",
}


def category_label(code: str) -> str:
    name = CATEGORY_LABELS.get(code, "")
    return f"Categoría {code} — {name}" if name else f"Categoría {code}"


def allocate_section_display_caps(
    section_cfg: Dict[str, Any], item_names: List[str]
) -> Dict[str, int]:
    return allocate_section_item_caps(section_cfg, item_names)


def audit_only_items(criteria: Dict[str, Any]) -> set:
    keys = set()
    for sec_name, cfg in criteria.get("sections", {}).items():
        for item_name, item in cfg.get("items", {}).items():
            if item.get("audit_only"):
                keys.add((sec_name, item_name))
    return keys


def filter_audit_items(
    df_items: pd.DataFrame, criteria: Dict[str, Any], include_audit: bool = False
) -> pd.DataFrame:
    if include_audit or df_items.empty:
        return df_items
    audit = audit_only_items(criteria)
    if not audit:
        return df_items
    mask = ~df_items.apply(lambda r: (r["Sección"], r["Ítem"]) in audit, axis=1)
    return df_items[mask].copy()


def results_to_dataframe(item_results, criteria: Optional[Dict[str, Any]] = None) -> pd.DataFrame:
    display_caps: Dict[tuple, Optional[int]] = {}
    if criteria:
        for sec_name, cfg in criteria.get("sections", {}).items():
            names = list(cfg.get("items", {}).keys())
            for item_name, cap in allocate_section_item_caps(cfg, names).items():
                display_caps[(sec_name, item_name)] = cap

    rows = []
    for r in item_results:
        applied = int(r.capped_item_points)
        tope = display_caps.get((r.section, r.item))
        if tope is None and criteria:
            cfg = criteria.get("sections", {}).get(r.section, {})
            item_cfg = cfg.get("items", {}).get(r.item, {})
            if float(item_cfg.get("max_points", r.item_max_points)) < 0:
                # Ítem sin tope fijo: mostrar el puntaje aplicado (evita "—").
                tope_display: Any = applied
            else:
                tope_display = int(r.item_max_points) if r.item_max_points >= 0 else applied
        elif tope is None:
            tope_display = applied if r.item_max_points < 0 else int(r.item_max_points)
        else:
            tope_display = int(tope)
        rows.append(
            {
                "Sección": r.section,
                "Ítem": r.item,
                "Ocurrencias": r.count,
                "Puntos unitarios": r.unit_points,
                "Puntaje bruto": r.raw_points,
                "Tope en sección": tope_display,
                "Puntaje (tope aplicado)": applied,
                "Evidencia (1er match)": r.evidence,
            }
        )
    return pd.DataFrame(rows)


def section_totals_dataframe(section_totals: Dict[str, float]) -> pd.DataFrame:
    df = pd.DataFrame([{"Sección": k, "Subtotal": v} for k, v in section_totals.items()])
    return df.sort_values("Subtotal", ascending=False)


def scoring_section_names(criteria: Dict[str, Any]) -> List[str]:
    return [
        name
        for name, cfg in criteria.get("sections", {}).items()
        if float(cfg.get("max_points", 0)) > 0
    ]


def section_totals_from_items(
    df_items: pd.DataFrame,
    criteria: Dict[str, Any],
    include_audit: bool = False,
) -> pd.DataFrame:
    """Subtotales por sección = suma de ítems, limitada por el máx. del apartado."""
    df = filter_audit_items(df_items, criteria, include_audit=include_audit)
    rows = []
    for name in scoring_section_names(criteria):
        sub = int(df.loc[df["Sección"] == name, "Puntaje (tope aplicado)"].sum())
        sec_max = int(round(float(criteria["sections"][name].get("max_points", 0))))
        if sec_max > 0:
            sub = min(sub, sec_max)
        rows.append({"Sección": name, "Subtotal": sub})
    return pd.DataFrame(rows).sort_values("Subtotal", ascending=False)


def export_excel(
    df_items: pd.DataFrame,
    df_sec_tot: pd.DataFrame,
    total: float,
    category: str,
    criteria: Dict[str, Any],
) -> bytes:
    excel_out = io.BytesIO()
    with pd.ExcelWriter(excel_out, engine="xlsxwriter") as writer:
        for section_name in criteria.get("sections", {}).keys():
            df_s = filter_audit_items(df_items, criteria)
            df_s = df_s[df_s["Sección"] == section_name].copy()
            if df_s.empty:
                continue
            df_s.to_excel(writer, sheet_name=section_name[:31], index=False)

        resumen = df_sec_tot.copy()
        resumen.loc[len(resumen)] = ["TOTAL", total]
        resumen.loc[len(resumen)] = ["CATEGORÍA", category_label(category)]
        resumen.to_excel(writer, sheet_name="RESUMEN", index=False)

    excel_out.seek(0)
    return excel_out.getvalue()


def export_word(
    df_items: pd.DataFrame,
    df_sec_tot: pd.DataFrame,
    total: float,
    category: str,
    cat_desc: str,
    filename: str,
    criteria: Dict[str, Any],
    include_evidence: bool = False,
) -> bytes:
    doc = Document()
    title = doc.add_paragraph("Universidad Católica de Cuyo")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Secretaría de Investigación — Categorización de Investigadores")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Informe de valoración de CVar (Anexo VII)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    doc.add_paragraph(f"Archivo evaluado: {filename}")
    doc.add_paragraph(f"Puntaje total: {total:.1f}")
    doc.add_paragraph(f"Categoría alcanzada: {category_label(category)}")
    if cat_desc:
        doc.add_paragraph(cat_desc)

    doc.add_paragraph("")
    doc.add_heading("Totales por sección", level=2)
    for _, row in df_sec_tot.iterrows():
        doc.add_paragraph(f"- {row['Sección']}: {float(row['Subtotal']):.1f}")

    scoring_sections = [
        name
        for name, cfg in criteria.get("sections", {}).items()
        if float(cfg.get("max_points", 0)) > 0
    ]

    for section_name in scoring_sections:
        doc.add_heading(section_name, level=2)
        df_s = filter_audit_items(df_items, criteria)
        df_s = df_s[df_s["Sección"] == section_name].copy()
        cols = ["Ítem", "Ocurrencias", "Puntaje (tope aplicado)", "Tope en sección"]
        if include_evidence:
            cols.append("Evidencia (1er match)")

        if df_s.empty:
            doc.add_paragraph("Sin ítems detectados.")
            continue

        tbl = doc.add_table(rows=1, cols=len(cols))
        hdr = tbl.rows[0].cells
        for i, col in enumerate(cols):
            hdr[i].text = col

        for _, r in df_s.iterrows():
            cells = tbl.add_row().cells
            for i, col in enumerate(cols):
                cells[i].text = str(r.get(col, ""))

        sec_rows = df_sec_tot[df_sec_tot["Sección"] == section_name]
        if not sec_rows.empty:
            doc.add_paragraph(f"Subtotal sección: {float(sec_rows['Subtotal'].values[0]):.1f}")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
